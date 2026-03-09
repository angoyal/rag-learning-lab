"""Document readers for various file formats."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

import bs4
import markdown
import pdfplumber
from docx import Document as DocxDocument


@dataclass
class Document:
    """A parsed document with extracted text and metadata.

    Attributes:
        text: Extracted text content.
        source: File path or name.
        file_type: One of "pdf", "docx", "markdown", "html", "txt".
        metadata: Format-specific metadata (e.g. page_count for PDFs).
    """

    text: str
    source: str
    file_type: str
    metadata: dict = field(default_factory=dict)


def read_text(path: Path) -> Document:
    """Read a plain text file into a Document.

    Args:
        path: Path to the .txt file.

    Returns:
        Document with the file's text content.

    Raises:
        FileNotFoundError: If the file does not exist.
    """
    path = Path(path)
    text = path.read_text(encoding="utf-8")
    return Document(text=text, source=str(path), file_type="txt")


def read_markdown(path: Path) -> Document:
    """Read a Markdown file, converting to plain text.

    Converts Markdown to HTML via the `markdown` library, then strips
    tags using BeautifulSoup to produce clean text.

    Args:
        path: Path to the .md file.

    Returns:
        Document with plain text extracted from the Markdown.
    """
    path = Path(path)
    raw = path.read_text(encoding="utf-8")
    html = markdown.markdown(raw, extensions=["fenced_code"])
    text = _extract_text_from_html(html)
    return Document(text=text, source=str(path), file_type="markdown")


def read_html(path: Path) -> Document:
    """Read an HTML file, extracting visible text.

    Strips script, style, nav, header, and footer tags before
    extracting text content.

    Args:
        path: Path to the .html or .htm file.

    Returns:
        Document with visible text extracted from the HTML.
    """
    path = Path(path)
    raw = path.read_text(encoding="utf-8")
    text = _extract_text_from_html(raw)
    return Document(text=text, source=str(path), file_type="html")


def read_pdf(path: Path) -> Document:
    """Read a PDF file, extracting text page by page.

    Pages are joined with double newlines so downstream chunkers
    can split on paragraph boundaries.

    Args:
        path: Path to the .pdf file.

    Returns:
        Document with extracted text and page_count in metadata.
    """
    path = Path(path)
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text)
    text = "\n\n".join(pages)
    return Document(
        text=text,
        source=str(path),
        file_type="pdf",
        metadata={"page_count": len(pages)},
    )


def read_docx(path: Path) -> Document:
    """Read a DOCX file, extracting paragraph text.

    Paragraphs are joined with double newlines.

    Args:
        path: Path to the .docx file.

    Returns:
        Document with extracted paragraph text.
    """
    path = Path(path)
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    text = "\n\n".join(paragraphs)
    return Document(text=text, source=str(path), file_type="docx")


_EXTENSION_MAP = {
    ".txt": read_text,
    ".md": read_markdown,
    ".html": read_html,
    ".htm": read_html,
    ".pdf": read_pdf,
    ".docx": read_docx,
}


def read_document(path: Path) -> Document:
    """Read a document file, auto-detecting format from its extension.

    Supports .txt, .md, .html, .htm, .pdf, and .docx files.

    Args:
        path: Path to the document file.

    Returns:
        Document with extracted text and metadata.

    Raises:
        ValueError: If the file extension is not supported.
    """
    path = Path(path)
    ext = path.suffix.lower()
    reader = _EXTENSION_MAP.get(ext)
    if reader is None:
        raise ValueError(f"Unsupported file extension: {ext!r}")
    return reader(path)


def _extract_text_from_html(html: str) -> str:
    """Extract visible text from HTML, stripping non-content tags."""
    soup = bs4.BeautifulSoup(html, "html.parser")
    for tag in soup.find_all(["script", "style", "nav", "header", "footer"]):
        tag.decompose()
    return soup.get_text(separator="\n").strip()
